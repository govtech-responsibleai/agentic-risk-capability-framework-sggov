"""Export utilities for generating Word documents."""

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, Any, List
# Import will be done inside the function to avoid relative import issues


def export_assessment_to_word() -> Document:
    """Export the complete risk assessment to a Word document with error handling.
    
    Returns:
        Document object or None if creation fails
    """
    try:
        doc = Document()
    except Exception as e:
        st.error(f"Failed to create Word document: {str(e)}")
        return None
    
    # Title
    title = doc.add_heading('Agentic Risk Capability Framework Assessment', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")
    
    # 1. Application Information
    doc.add_heading('1. Application Information', level=1)
    
    if 'application_description' in st.session_state:
        doc.add_heading('1.1 Generated Application Description', level=2)
        doc.add_paragraph(st.session_state.application_description)
    
    if 'application_info' in st.session_state:
        doc.add_heading('1.2 Application Summary', level=2)
        app_info = st.session_state.application_info
        doc.add_paragraph(f"Description: {app_info.get('description', 'Not provided')}")
        doc.add_paragraph(f"Data Classification: {app_info.get('data_classification', 'Not provided')}")
        doc.add_paragraph(f"Human in the Loop: {app_info.get('human_in_loop', 'Not provided')}")
        doc.add_paragraph(f"Public Facing: {app_info.get('public_facing', 'Not provided')}")
        doc.add_paragraph(f"Criticality: {app_info.get('criticality', 'Not provided')}")
        doc.add_paragraph(f"PII Data: {app_info.get('pii_data', 'Not provided')}")
        doc.add_paragraph(f"Components: {app_info.get('components', 'Not provided')}")
    
    # 2. Capability Analysis
    if 'capability_analysis' in st.session_state:
        doc.add_heading('2. System Capabilities Analysis', level=1)
        analysis = st.session_state.capability_analysis
        doc.add_paragraph(f"Analysis Reasoning: {analysis.reasoning}")
        
        doc.add_heading('2.1 Selected Applicable Capabilities', level=2)
        # Import here to avoid relative import issues
        from utils.data_loader import load_data
        capabilities, _, _, _, _ = load_data()
        for cap_id in analysis.applicable_capabilities:
            if cap_id in capabilities:
                cap_data = capabilities[cap_id]
                doc.add_paragraph(f"• {cap_id}: {cap_data['name']} ({cap_data['category']})")
    
    # 3. Risk Assessment
    if 'risk_assessments' in st.session_state and 'applicable_risks' in st.session_state:
        doc.add_heading('3. Risk Assessment', level=1)
        
        # Import here to avoid relative import issues
        from utils.data_loader import load_data
        
        # Inline get_controls_for_risk function to avoid import issues
        def get_controls_for_risk(risk_id: str, risks: Dict[str, Any], controls: Dict[str, Any]) -> List[Dict[str, Any]]:
            """Get controls for a specific risk with error handling."""
            risk_controls = []
            try:
                if not risks:
                    return []
                if risk_id in risks:
                    control_ids = risks[risk_id].get('controls', [])
                    for ctrl_id in control_ids:
                        if ctrl_id in controls:
                            risk_controls.append({
                                'id': ctrl_id,
                                'name': controls[ctrl_id]['name'],
                                'description': controls[ctrl_id]['description']
                            })
                return risk_controls
            except Exception:
                return []
        
        # Load data
        _, risks, controls, components, design = load_data()
        
        # Capability-specific risks
        doc.add_heading('3.1 Capability-Specific Risks', level=2)
        for risk_id in st.session_state.applicable_risks:
            if risk_id in risks:
                risk_data = risks[risk_id]
                if risk_data.get('capabilities') and not (risk_data.get('components') or risk_data.get('design')):
                    if risk_id in st.session_state.risk_assessments:
                        assessment = st.session_state.risk_assessments[risk_id]
                        
                        doc.add_heading(f"{risk_id}: {risk_data['name']}", level=3)
                        doc.add_paragraph(f"Description: {risk_data['description']}")
                        doc.add_paragraph(f"Contextualization: {assessment.context}")
                        
                        # Likelihood
                        doc.add_paragraph(f"Likelihood Score: {assessment.likelihood.score}/5")
                        doc.add_paragraph(f"Likelihood Reasoning: {assessment.likelihood.reasoning}")
                        
                        # Impact
                        doc.add_paragraph(f"Impact Score: {assessment.impact.score}/5")
                        doc.add_paragraph(f"Impact Reasoning: {assessment.impact.reasoning}")
                        doc.add_paragraph("")
        
        # Component and Design risks
        doc.add_heading('3.2 Component and Design Risks', level=2)
        for risk_id in st.session_state.applicable_risks:
            if risk_id in risks:
                risk_data = risks[risk_id]
                if (risk_data.get('components') or risk_data.get('design')) and not risk_data.get('capabilities'):
                    if risk_id in st.session_state.risk_assessments:
                        assessment = st.session_state.risk_assessments[risk_id]
                        
                        doc.add_heading(f"{risk_id}: {risk_data['name']}", level=3)
                        doc.add_paragraph(f"Description: {risk_data['description']}")
                        doc.add_paragraph(f"Contextualization: {assessment.context}")
                        
                        # Likelihood
                        doc.add_paragraph(f"Likelihood Score: {assessment.likelihood.score}/5")
                        doc.add_paragraph(f"Likelihood Reasoning: {assessment.likelihood.reasoning}")
                        
                        # Impact
                        doc.add_paragraph(f"Impact Score: {assessment.impact.score}/5")
                        doc.add_paragraph(f"Impact Reasoning: {assessment.impact.reasoning}")
                        doc.add_paragraph("")
    
    # 4. Controls
    if 'high_priority_risks' in st.session_state:
        doc.add_heading('4. Controls and Implementation', level=1)
        
        # Thresholds
        likelihood_threshold = st.session_state.get('likelihood_threshold', 4)
        impact_threshold = st.session_state.get('impact_threshold', 4)
        doc.add_paragraph(f"Control Thresholds: Likelihood ≥ {likelihood_threshold} AND Impact ≥ {impact_threshold}")
        doc.add_paragraph(f"High-Priority Risks: {len(st.session_state.high_priority_risks)} risks meet the threshold criteria")
        doc.add_paragraph("")
        
        # Controls for each high-priority risk
        for i, risk_id in enumerate(st.session_state.high_priority_risks, 1):
            if risk_id in risks:
                risk_data = risks[risk_id]
                doc.add_heading(f"4.{i} {risk_id}: {risk_data['name']}", level=2)
                
                # Risk assessment summary
                if risk_id in st.session_state.risk_assessments:
                    assessment = st.session_state.risk_assessments[risk_id]
                    doc.add_paragraph(f"Likelihood: {assessment.likelihood.score}/5, Impact: {assessment.impact.score}/5")
                    doc.add_paragraph(f"Context: {assessment.context}")
                    doc.add_paragraph("")
                
                # Controls
                risk_controls = get_controls_for_risk(risk_id, risks, controls)
                if risk_controls:
                    doc.add_paragraph("Controls:")
                    for j, control in enumerate(risk_controls, 1):
                        doc.add_paragraph(f"{j}. {control['id']}: {control['name']}")
                        doc.add_paragraph(f"   Description: {control['description']}")
                        
                        # Implementation status
                        control_key = f"control_implementation_{risk_id}_{control['id']}"
                        if control_key in st.session_state:
                            doc.add_paragraph(f"   Implementation Status: {st.session_state[control_key]}")
                        else:
                            doc.add_paragraph(f"   Implementation Status: I did not implement this control. I accept all residual risk.")
                        doc.add_paragraph("")
    
    return doc
